"""
Document generation service for creating DOCX and text files from templates.
"""

import io
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from jinja2 import Template as Jinja2Template

from src.core.logging import get_logger
from src.models.template import Template

logger = get_logger(__name__)


class DocumentGeneratorService:
    """Service for generating formatted documents from templates."""

    def __init__(self):
        """Initialize document generator service."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, DOCX generation will fail")
        logger.info("DocumentGeneratorService initialized")

    async def generate_document(
        self,
        template: Template,
        content: Dict[str, Any],
        output_path: Optional[Path] = None
    ) -> bytes:
        """Generate document from template and content.

        Args:
            template: Template configuration
            content: Content to populate template
            output_path: Optional path to save document

        Returns:
            Document bytes

        Raises:
            ValueError: If output format is unsupported
        """
        if template.output_format == "docx":
            return await self.generate_docx(template, content, output_path)
        elif template.output_format in ["txt", "md"]:
            return await self.generate_text(template, content, output_path)
        else:
            raise ValueError(f"Unsupported output format: {template.output_format}")

    async def generate_docx(
        self,
        template: Template,
        content: Dict[str, Any],
        output_path: Optional[Path] = None
    ) -> bytes:
        """Generate DOCX document.

        Args:
            template: Template configuration
            content: Content fields
            output_path: Optional path to save

        Returns:
            DOCX file bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation")

        # Check if template file exists
        template_path = Path(template.file_path)

        if template_path.exists() and template_path.suffix == ".dotx":
            # Load from Word template
            doc = Document(str(template_path))
        else:
            # Create new document
            doc = Document()

        # Populate based on template type
        if template.type == "general_meeting":
            self._populate_general_meeting(doc, content)
        elif template.type == "project_meeting":
            self._populate_project_meeting(doc, content)
        else:
            self._populate_generic(doc, content)

        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_bytes = doc_io.getvalue()

        # Optionally save to file
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(doc_bytes)
            logger.info(f"DOCX saved to {output_path}")

        return doc_bytes

    async def generate_text(
        self,
        template: Template,
        content: Dict[str, Any],
        output_path: Optional[Path] = None
    ) -> bytes:
        """Generate text document.

        Args:
            template: Template configuration
            content: Content fields
            output_path: Optional path to save

        Returns:
            Text file bytes
        """
        # Check for template file
        template_path = Path(template.file_path)

        if template_path.exists():
            # Load template file
            with open(template_path, "r") as f:
                template_text = f.read()

            # Use Jinja2 for templating
            jinja_template = Jinja2Template(template_text)
            text = jinja_template.render(**content)
        else:
            # Generate from content
            if template.type == "system_request":
                text = self._generate_system_request_text(content)
            else:
                text = self._generate_generic_text(content)

        text_bytes = text.encode("utf-8")

        # Optionally save to file
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w") as f:
                f.write(text)
            logger.info(f"Text saved to {output_path}")

        return text_bytes

    def _populate_general_meeting(self, doc: "Document", content: Dict[str, Any]) -> None:
        """Populate general meeting notes in DOCX.

        Args:
            doc: python-docx Document
            content: Content fields
        """
        # Title
        title = doc.add_heading("Meeting Notes", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        doc.add_heading("Meeting Details", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = "Light Grid Accent 1"

        table.rows[0].cells[0].text = "Date:"
        table.rows[0].cells[1].text = content.get("meeting_date", "N/A")
        table.rows[1].cells[0].text = "Attendees:"
        table.rows[1].cells[1].text = content.get("attendees", "N/A")

        # Discussion Topics
        doc.add_heading("Discussion Topics", level=2)
        doc.add_paragraph(content.get("discussion_topics", "No topics recorded"))

        # Decisions
        doc.add_heading("Key Decisions", level=2)
        decisions = content.get("decisions", "No decisions recorded")
        if isinstance(decisions, list):
            for decision in decisions:
                doc.add_paragraph(decision, style="List Bullet")
        else:
            doc.add_paragraph(decisions)

        # Action Items
        doc.add_heading("Action Items", level=2)
        action_items = content.get("action_items", "No action items")
        if isinstance(action_items, list):
            for item in action_items:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(action_items)

        # Next Steps
        doc.add_heading("Next Steps", level=2)
        doc.add_paragraph(content.get("next_steps", "No next steps defined"))

    def _populate_project_meeting(self, doc: "Document", content: Dict[str, Any]) -> None:
        """Populate project meeting notes in DOCX.

        Args:
            doc: python-docx Document
            content: Content fields
        """
        # Title
        project_name = content.get("project_name", "Project")
        title = doc.add_heading(f"{project_name} - Meeting Notes", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        doc.add_heading("Meeting Details", level=2)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Light Grid Accent 1"

        table.rows[0].cells[0].text = "Project:"
        table.rows[0].cells[1].text = project_name
        table.rows[1].cells[0].text = "Date:"
        table.rows[1].cells[1].text = content.get("meeting_date", "N/A")
        table.rows[2].cells[0].text = "Attendees:"
        table.rows[2].cells[1].text = content.get("attendees", "N/A")

        # Status Updates
        doc.add_heading("Status Updates", level=2)
        doc.add_paragraph(content.get("status_updates", "No updates provided"))

        # Risks & Issues
        doc.add_heading("Risks & Issues", level=2)
        risks = content.get("risks_issues", "No risks or issues reported")
        doc.add_paragraph(risks)

        # Decisions
        doc.add_heading("Decisions Made", level=2)
        decisions = content.get("decisions", "No decisions recorded")
        if isinstance(decisions, list):
            for decision in decisions:
                doc.add_paragraph(decision, style="List Bullet")
        else:
            doc.add_paragraph(decisions)

        # Action Items
        doc.add_heading("Action Items", level=2)
        action_items = content.get("action_items", "No action items")
        if isinstance(action_items, list):
            table = doc.add_table(rows=len(action_items) + 1, cols=3)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Item"
            table.rows[0].cells[1].text = "Owner"
            table.rows[0].cells[2].text = "Due Date"

            for idx, item in enumerate(action_items, 1):
                if isinstance(item, dict):
                    table.rows[idx].cells[0].text = item.get("description", "")
                    table.rows[idx].cells[1].text = item.get("owner", "")
                    table.rows[idx].cells[2].text = item.get("due_date", "")
                else:
                    table.rows[idx].cells[0].text = str(item)
        else:
            doc.add_paragraph(action_items)

        # Next Meeting
        doc.add_heading("Next Meeting", level=2)
        doc.add_paragraph(content.get("next_meeting", "TBD"))

    def _populate_generic(self, doc: "Document", content: Dict[str, Any]) -> None:
        """Populate generic document.

        Args:
            doc: python-docx Document
            content: Content fields
        """
        for key, value in content.items():
            doc.add_heading(key.replace("_", " ").title(), level=2)
            doc.add_paragraph(str(value))

    def _generate_system_request_text(self, content: Dict[str, Any]) -> str:
        """Generate system request text format.

        Args:
            content: Content fields

        Returns:
            Formatted text
        """
        feature_name = content.get("feature_name", "Untitled Feature")
        text = f"""# Feature Request: {feature_name}

**Date:** {datetime.utcnow().strftime("%Y-%m-%d")}

## Problem Statement

{content.get("problem_statement", "[Not provided]")}

## Target Users

{content.get("target_users", "[Not provided]")}

## Proposed Solution

{content.get("proposed_solution", "[Not provided]")}

## Acceptance Criteria

{content.get("acceptance_criteria", "[Not provided]")}

## Technical Constraints

{content.get("technical_constraints", "[Not provided]")}

## Dependencies

{content.get("dependencies", "[Not provided]")}

## Priority

{content.get("priority", "[Not provided]")}

## Additional Notes

{content.get("additional_notes", "[None]")}

---

This feature request was generated from transcript analysis.
For spec-kit processing, ensure all sections are complete.
"""
        return text

    def _generate_generic_text(self, content: Dict[str, Any]) -> str:
        """Generate generic text format.

        Args:
            content: Content fields

        Returns:
            Formatted text
        """
        lines = [f"# Document\n\n**Generated:** {datetime.utcnow().isoformat()}\n"]

        for key, value in content.items():
            lines.append(f"\n## {key.replace('_', ' ').title()}\n")
            lines.append(f"{value}\n")

        return "\n".join(lines)
